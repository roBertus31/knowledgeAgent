"""
file_parser.py
Extracts plain text from .docx, .pdf, .txt, .md, and other text-like files.
Returns (text: str, metadata: dict) tuples.
"""

from __future__ import annotations

import re
import sys
from pathlib import Path


# ── helpers ──────────────────────────────────────────────────────────────────

def _clean(text: str) -> str:
    """Collapse excessive blank lines; strip trailing whitespace per line."""
    lines = [line.rstrip() for line in text.splitlines()]
    # collapse 3+ consecutive blank lines to 2
    result, blank_count = [], 0
    for line in lines:
        if line == "":
            blank_count += 1
            if blank_count <= 2:
                result.append(line)
        else:
            blank_count = 0
            result.append(line)
    return "\n".join(result).strip()


# ── format handlers ───────────────────────────────────────────────────────────

def _parse_txt(path: Path) -> tuple[str, dict]:
    try:
        text = path.read_text(encoding="utf-8", errors="replace")
    except Exception as e:
        return "", {"error": str(e)}
    return _clean(text), {"encoding": "utf-8"}


def _parse_md(path: Path) -> tuple[str, dict]:
    return _parse_txt(path)


def _parse_docx(path: Path) -> tuple[str, dict]:
    try:
        import docx  # python-docx
    except ImportError:
        return "", {"error": "python-docx not installed. Run: pip install python-docx"}

    try:
        doc = docx.Document(str(path))
    except Exception as e:
        return "", {"error": f"Could not open docx: {e}"}

    parts = []

    # Core paragraph text (preserves heading levels as markdown)
    for para in doc.paragraphs:
        style = para.style.name if para.style else ""
        text = para.text.strip()
        if not text:
            parts.append("")
            continue
        if style.startswith("Heading 1"):
            parts.append(f"# {text}")
        elif style.startswith("Heading 2"):
            parts.append(f"## {text}")
        elif style.startswith("Heading 3"):
            parts.append(f"### {text}")
        elif style.startswith("Heading"):
            parts.append(f"#### {text}")
        else:
            parts.append(text)

    # Tables
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
        if rows:
            # Insert a markdown table separator after the first row
            sep = "| " + " | ".join(["---"] * len(table.rows[0].cells)) + " |"
            rows.insert(1, sep)
            parts.append("")
            parts.extend(rows)
            parts.append("")

    meta = {
        "core_properties": {},
    }
    try:
        cp = doc.core_properties
        meta["core_properties"] = {
            "author": cp.author or "",
            "created": str(cp.created) if cp.created else "",
            "modified": str(cp.modified) if cp.modified else "",
            "title": cp.title or "",
            "subject": cp.subject or "",
        }
    except Exception:
        pass

    return _clean("\n".join(parts)), meta


def _parse_pdf(path: Path) -> tuple[str, dict]:
    # Try pymupdf first (fitz) — better layout preservation
    try:
        import fitz  # pymupdf
        doc = fitz.open(str(path))
        pages = []
        for page_num, page in enumerate(doc, start=1):
            page_text = page.get_text("text")
            if page_text.strip():
                pages.append(f"[Page {page_num}]\n{page_text.strip()}")
        full_text = "\n\n".join(pages)
        meta = {
            "page_count": len(doc),
            "metadata": doc.metadata,
        }
        doc.close()
        return _clean(full_text), meta
    except ImportError:
        pass
    except Exception as e:
        pass  # fall through to pdfplumber

    # Fallback: pdfplumber
    try:
        import pdfplumber
        pages = []
        meta = {}
        with pdfplumber.open(str(path)) as pdf:
            meta["page_count"] = len(pdf.pages)
            for page_num, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ""
                if text.strip():
                    pages.append(f"[Page {page_num}]\n{text.strip()}")
        return _clean("\n\n".join(pages)), meta
    except ImportError:
        return "", {"error": "Neither pymupdf nor pdfplumber is installed."}
    except Exception as e:
        return "", {"error": str(e)}


def _parse_csv(path: Path) -> tuple[str, dict]:
    """Convert CSV to a markdown-style representation."""
    import csv
    try:
        rows = []
        with path.open(encoding="utf-8", errors="replace", newline="") as f:
            reader = csv.reader(f)
            for row in reader:
                rows.append(row)
        if not rows:
            return "", {}
        # Build markdown table
        lines = []
        lines.append("| " + " | ".join(rows[0]) + " |")
        lines.append("| " + " | ".join(["---"] * len(rows[0])) + " |")
        for row in rows[1:]:
            lines.append("| " + " | ".join(row) + " |")
        return "\n".join(lines), {"row_count": len(rows) - 1}
    except Exception as e:
        return "", {"error": str(e)}


# ── public API ────────────────────────────────────────────────────────────────

SUPPORTED_EXTENSIONS = {
    ".txt": _parse_txt,
    ".md":  _parse_md,
    ".docx": _parse_docx,
    ".pdf": _parse_pdf,
    ".csv": _parse_csv,
    # Additional plain-text formats
    ".log":  _parse_txt,
    ".rst":  _parse_txt,
    ".html": _parse_txt,  # raw HTML — won't strip tags but captures text
}


def parse_file(path: str | Path) -> tuple[str, dict]:
    """
    Extract text and metadata from a file.

    Returns:
        (text, metadata) — text is the extracted content as a string;
        metadata is a dict with format-specific info or an 'error' key.
    """
    path = Path(path)
    if not path.exists():
        return "", {"error": f"File not found: {path}"}

    suffix = path.suffix.lower()
    handler = SUPPORTED_EXTENSIONS.get(suffix)

    if handler is None:
        # Try reading as plain text anyway
        try:
            text = path.read_text(encoding="utf-8", errors="replace")
            return _clean(text), {"warning": f"Unknown extension {suffix!r}, treated as plain text"}
        except Exception as e:
            return "", {"error": f"Unsupported file type {suffix!r} and could not read as text: {e}"}

    return handler(path)


def supported_extensions() -> list[str]:
    return sorted(SUPPORTED_EXTENSIONS.keys())


# ── CLI quick-test ────────────────────────────────────────────────────────────

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(f"Usage: python file_parser.py <file>")
        print(f"Supported: {', '.join(supported_extensions())}")
        sys.exit(1)

    target = Path(sys.argv[1])
    text, meta = parse_file(target)
    print(f"=== Metadata: {meta}")
    print(f"=== Extracted text ({len(text)} chars) ===")
    print(text[:3000])
    if len(text) > 3000:
        print(f"\n... [{len(text) - 3000} more chars] ...")
